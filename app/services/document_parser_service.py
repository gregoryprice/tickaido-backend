#!/usr/bin/env python3
"""
DocumentParserService - Extract structured content from documents (PDFs, Word, Excel, etc.)
Similar to AWS Textract functionality but using local libraries.
"""

import logging
from io import BytesIO
from typing import Any, Dict, List

try:
    import pymupdf  # PyMuPDF library - correct import
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    import magic
    HAS_MAGIC = True
except ImportError:
    HAS_MAGIC = False

from app.config.settings import get_settings

logger = logging.getLogger(__name__)


class DocumentParserService:
    """Document structure parsing with layout detection"""
    
    def __init__(self):
        self.settings = get_settings()
        
    async def analyze_document(self, content: bytes, features: List[str]) -> Dict[str, Any]:
        """
        Analyze document with specified features (TEXT, TABLES, FORMS, LAYOUT)
        
        Args:
            content: Raw document bytes
            features: List of analysis features to perform
            
        Returns:
            Structured document analysis with pages, blocks, and metadata
        """
        mime_type = self._detect_mime_type(content)
        logger.debug(f"DocumentParserService detected MIME type: {mime_type} for {len(content)} bytes")
        
        if mime_type == "application/pdf":
            logger.debug("DEBUG: DocumentParserService using PDF analysis path")
            return await self._analyze_pdf(content, features)
        elif mime_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
            logger.debug("DEBUG: DocumentParserService using Word analysis path")
            return await self._analyze_word_doc(content, features)
        elif mime_type in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"]:
            logger.debug("DEBUG: DocumentParserService using Excel analysis path")
            return await self._analyze_excel_doc(content, features)
        elif mime_type.startswith("text/"):
            logger.info("DocumentParserService: Using text file analysis path")
            return await self._analyze_text_file(content, features)
        else:
            raise ValueError(f"Unsupported document type: {mime_type}")
    
    async def _analyze_pdf(self, content: bytes, features: List[str]) -> Dict[str, Any]:
        """Analyze PDF document with structure detection"""
        
        logger.info(f"DEBUG: _analyze_pdf called with features {features}")

        if not HAS_PYMUPDF:
            raise ValueError("PyMuPDF not available - install pymupdf package for PDF text extraction")
        
        doc = pymupdf.open(stream=content, filetype="pdf")
        pages = []
        
        for page in doc:  # Iterate directly over pages
            page_data = {
                "page_number": page.number + 1,  # PyMuPDF page numbers are 0-based
                "text": "",
                "blocks": [],
                "confidence": 0.95
            }
            
            if "TEXT" in features:
                try:
                    # Extract text exactly as shown in official PyMuPDF docs
                    text = page.get_text()  # get plain text (is in UTF-8)
                    logger.info(f"DEBUG: PyMuPDF returned text: {text}")

                    # Check if PyMuPDF returned PDF source code instead of readable text
                    # This happens when PyMuPDF reads a malformed PDF or PDF source code
                    if text and len(text) > 10:
                        # Check for obvious PDF source markers
                        is_pdf_source = (
                            text.startswith('%PDF') or 
                            text.startswith('1 0 obj') or
                            (text.count('obj') > 2 and text.count('endobj') > 2) or
                            ('stream' in text and 'endstream' in text) or
                            ('xref' in text and '%%EOF' in text)
                        )
                        
                        if is_pdf_source:
                            print("DEBUG: Detected PDF source code in extracted text, clearing...")
                            logger.warning("PyMuPDF returned PDF source code instead of readable text")
                            # Set empty text so it can be handled appropriately
                            text = ""
                    
                    page_data["text"] = text
                    
                except Exception as e:
                    # If extraction fails, log and set empty text
                    logger.warning(f"PyMuPDF text extraction failed for page {page.number + 1}: {e}")
                    page_data["text"] = ""
            
            if "LAYOUT" in features:
                # Get text blocks with positioning
                text_dict = page.get_text("dict")
                blocks = text_dict.get("blocks", [])
                for block in blocks:
                    if "lines" in block:  # Text block
                        block_text = ""
                        for line in block["lines"]:
                            for span in line["spans"]:
                                block_text += span.get("text", "")
                        
                        if block_text.strip():  # Only add non-empty blocks
                            page_data["blocks"].append({
                                "type": "paragraph",
                                "text": block_text,
                                "confidence": 0.95,
                                "geometry": {
                                    "bounding_box": {
                                        "left": block["bbox"][0] / page.rect.width,
                                        "top": block["bbox"][1] / page.rect.height,
                                        "width": (block["bbox"][2] - block["bbox"][0]) / page.rect.width,
                                        "height": (block["bbox"][3] - block["bbox"][1]) / page.rect.height
                                    }
                                }
                            })
            
            if "TABLES" in features:
                # Extract tables using layout analysis
                tables = await self._extract_pdf_tables(page)
                page_data["blocks"].extend(tables)
            
            pages.append(page_data)
        
        doc.close()
        
        return {
            "pages": pages,
            "metadata": {
                "total_pages": len(pages),
                "document_type": "pdf",
                "language": "auto-detected"
            }
        }
    
    async def _analyze_word_doc(self, content: bytes, features: List[str]) -> Dict[str, Any]:
        """Analyze Word document"""
        from docx import Document
        
        doc = Document(BytesIO(content))
        
        page_data = {
            "page_number": 1,
            "text": "",
            "blocks": [],
            "confidence": 0.98
        }
        
        if "TEXT" in features:
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            page_data["text"] = "\n".join(full_text)
        
        if "LAYOUT" in features:
            for para in doc.paragraphs:
                if para.text.strip():
                    page_data["blocks"].append({
                        "type": "paragraph",
                        "text": para.text,
                        "confidence": 0.98
                    })
        
        if "TABLES" in features:
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append({"cells": cells})
                
                page_data["blocks"].append({
                    "type": "table",
                    "rows": rows,
                    "confidence": 0.95
                })
        
        return {
            "pages": [page_data],
            "metadata": {
                "total_pages": 1,
                "document_type": "word",
                "language": "auto-detected"
            }
        }
    
    async def _analyze_excel_doc(self, content: bytes, features: List[str]) -> Dict[str, Any]:
        """Analyze Excel document"""
        import openpyxl
        
        workbook = openpyxl.load_workbook(BytesIO(content))
        pages = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            page_data = {
                "page_number": len(pages) + 1,
                "text": "",
                "blocks": [],
                "confidence": 0.98
            }
            
            if "TEXT" in features or "TABLES" in features:
                # Extract all cell values
                rows = []
                all_text = []
                
                for row in sheet.iter_rows(values_only=True):
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    if any(val.strip() for val in row_values):  # Skip empty rows
                        rows.append({"cells": row_values})
                        all_text.extend([val for val in row_values if val.strip()])
                
                if "TEXT" in features:
                    page_data["text"] = " ".join(all_text)
                
                if "TABLES" in features and rows:
                    page_data["blocks"].append({
                        "type": "table",
                        "rows": rows,
                        "confidence": 0.98
                    })
            
            pages.append(page_data)
        
        workbook.close()
        
        return {
            "pages": pages,
            "metadata": {
                "total_pages": len(pages),
                "document_type": "excel",
                "language": "auto-detected"
            }
        }
    
    async def _analyze_text_file(self, content: bytes, features: List[str]) -> Dict[str, Any]:
        """Analyze plain text file"""
        # Validate that this is actually readable text
        try:
            text_content = content.decode('utf-8')
            # Check if content seems like readable text
            if len(text_content.strip()) < 5:
                raise ValueError("Content too short to be valid document")
        except UnicodeDecodeError:
            try:
                text_content = content.decode('latin-1')
            except UnicodeDecodeError:
                raise ValueError("Cannot decode text content - corrupted or binary file")
        
        page_data = {
            "page_number": 1,
            "text": "",
            "blocks": [],
            "confidence": 0.99
        }
        
        if "TEXT" in features:
            page_data["text"] = text_content
        
        if "LAYOUT" in features:
            # Split into paragraphs
            paragraphs = text_content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    page_data["blocks"].append({
                        "type": "paragraph",
                        "text": para.strip(),
                        "confidence": 0.99
                    })
        
        return {
            "pages": [page_data],
            "metadata": {
                "total_pages": 1,
                "document_type": "text",
                "language": "auto-detected"
            }
        }
    
    async def _extract_pdf_tables(self, page) -> List[Dict[str, Any]]:
        """Extract tables from PDF page using basic layout analysis"""
        # Simple table detection - can be enhanced with more sophisticated algorithms
        tables = []
        
        # For now, return empty list - table extraction is complex
        # In production, could use libraries like tabula-py or camelot
        
        return tables
    
    async def extract_text(self, content: bytes) -> str:
        """Simple text extraction without structure"""
        analysis = await self.analyze_document(content, ["TEXT"])
        return "\n".join([page["text"] for page in analysis["pages"]])
    
    def _detect_mime_type(self, content: bytes) -> str:
        """Detect document MIME type from content"""
        print(f"DEBUG: _detect_mime_type called with {len(content)} bytes, HAS_MAGIC={HAS_MAGIC}")
        print(f"DEBUG: Content starts with: {repr(content[:20])}")
        
        if HAS_MAGIC:
            result = magic.from_buffer(content, mime=True)
            print(f"DEBUG: python-magic detected: {result}")
            return result
        
        # Fallback MIME type detection when python-magic isn't available
        if content.startswith(b'%PDF'):
            print("DEBUG: Fallback detection: application/pdf")
            return "application/pdf"
        elif content.startswith(b'PK\x03\x04'):
            # ZIP-based formats (Office documents)
            # Could be more specific but this is safer
            if b'word/' in content[:1024]:
                return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            elif b'xl/' in content[:1024]:
                return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            else:
                return "application/zip"
        elif content.startswith(b'\xFF\xFE') or content.startswith(b'\xFE\xFF'):
            return "text/plain"  # Unicode text
        else:
            # Check if content appears to be text
            try:
                decoded = content[:1024].decode('utf-8')  # Check first 1KB only
                if len(decoded.strip()) > 5:
                    return "text/plain"
                else:
                    return "application/octet-stream"
            except UnicodeDecodeError:
                return "application/octet-stream"